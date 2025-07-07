import tempfile
import json
import hashlib
import redis
from typing import List, Tuple, Dict, Any, Optional
import httpx
import chromadb
import io
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.chat_models import ChatOllama
from langchain_community.embeddings import OllamaEmbeddings
from langchain_chroma import Chroma
from langchain_groq import ChatGroq
from langchain.prompts import ChatPromptTemplate
from chromadb.config import Settings as ChromaSettings

from app.core.config import settings
from app.services import storage_service
from app.db.models import Project, User
import logging

logger = logging.getLogger(__name__)

class RAGService:

    def _ensure_ollama_model_is_available(self, model_name: str) -> None:
        if not settings.OLLAMA_HOST:
            return
        try:
            client = httpx.Client(base_url=settings.OLLAMA_HOST, timeout=30.0)
            response = client.get("/api/tags")
            response.raise_for_status()
            models = response.json().get("models", [])
            model_exists = any(m['name'].split(':')[0] == model_name.split(':')[0] for m in models)

            if model_exists:
                logger.info(f"Ollama model '{model_name}' is already available.")
                return

            logger.info(f"Ollama model '{model_name}' not found. Pulling it now. This may take a while...")
            pull_response = client.post("/api/pull", json={"name": model_name, "stream": False}, timeout=None)
            pull_response.raise_for_status()
            
            status_data = pull_response.json()
            if "error" in status_data:
                logger.error(f"Failed to pull Ollama model '{model_name}': {status_data['error']}")
                raise Exception(f"Failed to pull Ollama model: {status_data['error']}")
            
            logger.info(f"Successfully pulled Ollama model '{model_name}'.")

        except httpx.RequestError as e:
            logger.error(f"Error communicating with Ollama at {settings.OLLAMA_HOST}: {e}.")
            raise ConnectionError(f"Could not connect to Ollama service at {settings.OLLAMA_HOST}.")
        except Exception as e:
            logger.error(f"An unexpected error occurred while ensuring Ollama model '{model_name}' is available: {e}", exc_info=True)
            raise

    def __init__(self, user: User, project: Project) -> None:
        self.user: User = user
        self.project: Project = project
        self.collection_name: str = f"proj_{str(project.id).replace('-', '')}"

        if self.project.llm_provider == "ollama" and settings.OLLAMA_HOST:
            model_name = self.project.llm_model_name or settings.OLLAMA_MODEL
            logger.info(f"Using Ollama provider at {settings.OLLAMA_HOST} with model '{model_name}' for project '{self.project.name}'")
            self._ensure_ollama_model_is_available(model_name)
            self.llm = ChatOllama(base_url=settings.OLLAMA_HOST, model=model_name)
            self.embedding_function = OllamaEmbeddings(base_url=settings.OLLAMA_HOST, model=model_name)
        
        else:
            if self.project.llm_provider != 'groq':
                logger.warning(f"LLM provider '{self.project.llm_provider}' is not 'ollama' or Ollama is not configured. Defaulting to 'groq'.")
            
            model_name = self.project.llm_model_name or settings.LLM_MODEL_NAME
            logger.info(f"Using Groq provider with model '{model_name}' for project '{self.project.name}'")
            self.embedding_function = GoogleGenerativeAIEmbeddings(model=settings.EMBEDDING_MODEL_NAME)
            self.llm = ChatGroq(
                groq_api_key=settings.GROQ_API_KEY,
                model_name=model_name
            )

        try:
            self.redis_client: Optional[redis.Redis] = redis.from_url(settings.CELERY_BROKER_URL)
            self.redis_client.ping()
            logger.info("Successfully connected to Redis for caching.")
        except Exception as e:
            logger.error(f"Could not connect to Redis: {e}. Caching will be disabled.", exc_info=True)
            self.redis_client = None

        chroma_client = chromadb.PersistentClient(path=settings.CHROMA_PATH, settings=ChromaSettings(anonymized_telemetry=False))
        self.vectorstore: Chroma = Chroma(
            client=chroma_client, 
            collection_name=self.collection_name,
            embedding_function=self.embedding_function,
        )

    def _load_docs_from_file_obj(self, file_obj, file_type: str, file_name: str) -> list[LangchainDocument]:
        if file_type == "application/pdf":
            with tempfile.NamedTemporaryFile(delete=True, suffix=".pdf") as tmp:
                tmp.write(file_obj.read())
                tmp.seek(0)
                loader = PyPDFLoader(tmp.name)
                return loader.load()
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = DocxDocument(file_obj)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            return [LangchainDocument(page_content=full_text, metadata={"source": file_name})]
        elif file_type.startswith("text/"):
            content = file_obj.read().decode('utf-8', errors='ignore')
            return [LangchainDocument(page_content=content, metadata={"source": file_name})]
        else:
            logger.warning(f"Unsupported file type '{file_type}' for {file_name}. Skipping.")
            return []

    def process_document(
        self,
        storage_key: str,
        file_type: str,
        file_name: str,
        document_id: str,
        url: Optional[str] = None
    ) -> None:
        logger.info(f"Processing document_id: {document_id} for project {self.project.name}")
        docs: list[LangchainDocument] = []

        if url:
            try:
                response = httpx.get(url, follow_redirects=True, timeout=20.0)
                response.raise_for_status()
                soup = BeautifulSoup(response.content, 'html.parser')
                page_text = soup.get_text(separator='\n', strip=True)
                docs = [LangchainDocument(page_content=page_text, metadata={"source": url})]
            except Exception as e:
                logger.error(f"Failed to process URL {url}: {e}")
                return
        else:
            try:
                file_obj = io.BytesIO()
                storage_service.s3_client.download_fileobj(storage_service.BUCKET_NAME, storage_key, file_obj)
                file_obj.seek(0)
                docs = self._load_docs_from_file_obj(file_obj, file_type, file_name)
            except Exception as e:
                logger.error(f"Failed to download or process file {storage_key}: {e}")
                return

        if not docs:
            logger.warning(f"No content could be loaded from {file_name or url}. Skipping.")
            return

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE, chunk_overlap=settings.CHUNK_OVERLAP
        )
        chunks: list[LangchainDocument] = text_splitter.split_documents(docs)

        if not chunks:
            logger.warning(f"No text could be extracted from {file_name or url}. Skipping.")
            return

        for chunk in chunks:
            chunk.metadata['document_id'] = document_id
            chunk.metadata.setdefault('source', url or file_name)

        self.vectorstore.add_documents(documents=chunks)
        logger.info(f"Added {len(chunks)} chunks for document {document_id} to '{self.collection_name}'.")
        self.clear_cache_for_project()

    def delete_document_chunks(self, document_id: str) -> None:
        logger.info(f"Deleting chunks for document_id: {document_id} from '{self.collection_name}'")
        try:
            collection = self.vectorstore._collection
            existing_chunks = collection.get(where={"document_id": document_id}, include=[])
            if not existing_chunks or not existing_chunks.get('ids'):
                 logger.info(f"No chunks found for document_id: {document_id}. Nothing to delete.")
                 return
            collection.delete(where={"document_id": document_id})
            logger.info(f"Deleted {len(existing_chunks['ids'])} chunks for document_id: {document_id}.")
            self.clear_cache_for_project()
        except Exception as e:
            logger.error(f"Error deleting chunks for document {document_id}: {e}", exc_info=True)

    def clear_cache_for_project(self) -> None:
        if not self.redis_client:
            return
        try:
            keys_to_delete: List[str] = [k.decode('utf-8') for k in self.redis_client.scan_iter(f"rag_cache:{self.project.id}:*")]
            if keys_to_delete:
                self.redis_client.delete(*keys_to_delete)
                logger.info(f"Invalidated {len(keys_to_delete)} cache entries for project {self.project.id}.")
        except Exception as e:
            logger.error(f"Failed to clear Redis cache for project {self.project.id}: {e}", exc_info=True)

    def query(self, message: str) -> Tuple[str, List[Dict[str, Any]]]:
        if self.redis_client:
            message_hash: str = hashlib.sha256(message.encode()).hexdigest()
            cache_key: str = f"rag_cache:{self.project.id}:{message_hash}"
            try:
                cached_result = self.redis_client.get(cache_key)
                if cached_result:
                    logger.info(f"Returning cached RAG result for project {self.project.id}")
                    result = json.loads(cached_result)
                    return result['answer'], result['sources']
            except Exception as e:
                logger.warning(f"Redis cache check failed: {e}", exc_info=True)
        
        logger.info(f"Querying project '{self.project.name}' with: '{message}'")

        # **FIX**: Reverted to a faster, simpler retriever to improve performance.
        retriever = self.vectorstore.as_retriever(search_kwargs={"k": 7})
        
        try:
            if not self.vectorstore._collection or self.vectorstore._collection.count() == 0:
                 return "This project has no documents processed yet. Please upload a document and wait for it to complete.", []
        except Exception:
            return "This project has no documents processed yet. Please upload a document and wait for it to complete.", []
        
        # Use the simple retriever directly.
        final_docs: List[LangchainDocument] = retriever.get_relevant_documents(message)
        
        if not final_docs:
            return "I couldn't find relevant information in your documents to answer the query.", []

        context_text: str = "\n\n---\n\n".join([doc.page_content for doc in final_docs])
        
        rag_prompt = ChatPromptTemplate.from_template("""
            You are a helpful and diligent AI assistant. Your task is to answer the user's question based on the provided context.

            **Instructions:**
            1.  **Analyze the Context:** Carefully read the following context, which is composed of text chunks from one or more documents.
            2.  **Synthesize an Answer:** Formulate a comprehensive, well-structured answer to the user's question. Do not just copy-paste from the context. You must synthesize the information.
            3.  **Strictly Ground Your Answer:** Your entire response must be based *only* on the information available in the provided context. Do not use any outside knowledge.
            4.  **Handle Missing Information:** If the context does not contain the necessary information to answer the question, do not invent an answer. Instead, clearly state that the provided documents do not contain the answer.
            5.  **Be Conversational:** Present the answer in a clear and helpful manner.

            **Context:**
            {context}

            **Question:**
            {question}

            **Answer:**
        """)

        rag_chain = rag_prompt | self.llm
        answer: str = rag_chain.invoke({"context": context_text, "question": message}).content
        
        unique_sources: Dict[str, Dict[str, Any]] = {}
        for doc in final_docs:
            source_info: Dict[str, Any] = {"content": doc.page_content, "source": doc.metadata.get("source", "Unknown")}
            unique_sources[doc.page_content] = source_info
        
        sources: List[Dict[str, Any]] = list(unique_sources.values())
        
        if self.redis_client:
            try:
                result_to_cache: Dict[str, Any] = {"answer": answer, "sources": sources}
                self.redis_client.set(cache_key, json.dumps(result_to_cache), ex=3600)
                logger.info(f"Cached RAG result for project {self.project.id}")
            except Exception as e:
                logger.warning(f"Redis cache set failed: {e}", exc_info=True)

        return answer, sources